# -*- coding: utf-8 -*-
import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_report():
    doc = docx.Document()

    # Cấu hình Margins (lề trang): Cỡ chuẩn A4, lề 2.54cm (1 inch) cho tất cả các chiều
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.header_distance = Inches(0.5)
        section.footer_distance = Inches(0.5)

    # Thiết lập font mặc định cho toàn bộ tài liệu là Times New Roman
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.line_spacing = 1.3
    style.paragraph_format.space_after = Pt(6)

    # Hàm tiện ích để áp dụng font Times New Roman thủ công cho từng run (để đảm bảo không bị lỗi font của Word)
    def apply_run_font(run, size_pt=13, bold=False, italic=False, color_rgb=(0,0,0)):
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size_pt)
        run.bold = bold
        run.italic = italic
        run.font.color.rgb = RGBColor(*color_rgb)

    # Hàm viết tiêu đề cấp 1 (H1)
    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        apply_run_font(run, size_pt=16, bold=True)
        return p

    # Hàm viết tiêu đề cấp 2 (H2)
    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        apply_run_font(run, size_pt=14, bold=True)
        return p

    # Hàm viết tiêu đề cấp 3 (H3)
    def add_heading_3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        apply_run_font(run, size_pt=13, bold=True, italic=True)
        return p

    # Hàm thêm đoạn văn bản thường
    def add_body(text, bold_prefix="", indent_inches=0, italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(indent_inches)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if bold_prefix:
            r_pref = p.add_run(bold_prefix)
            apply_run_font(r_pref, size_pt=13, bold=True)
        r_text = p.add_run(text)
        apply_run_font(r_text, size_pt=13, italic=italic)
        return p

    # Hàm thêm bullet point
    def add_bullet(text, bold_prefix="", indent_inches=0.25):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(indent_inches)
        p.paragraph_format.space_after = Pt(4)
        if bold_prefix:
            r_pref = p.add_run(bold_prefix)
            apply_run_font(r_pref, size_pt=13, bold=True)
        r_text = p.add_run(text)
        apply_run_font(r_text, size_pt=13)
        return p

    # Hàm thêm hình ảnh
    def add_image(image_path, width_inches=5.0, caption=""):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run()
        try:
            r.add_picture(image_path, width=Inches(width_inches))
            if caption:
                p_cap = doc.add_paragraph()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_cap.paragraph_format.space_after = Pt(12)
                p_cap.paragraph_format.keep_with_next = True
                r_cap = p_cap.add_run(caption)
                apply_run_font(r_cap, size_pt=11, italic=True, color_rgb=(100, 100, 100))
        except Exception as e:
            print(f"Error adding picture {image_path}: {e}")
        return p

    # ────────────────────────────────────────────────────────
    # 1. TRANG BÌA (COVER PAGE)
    # ────────────────────────────────────────────────────────
    
    # Đại học Phenikaa
    p_uni = doc.add_paragraph()
    p_uni.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_uni.paragraph_format.space_before = Pt(20)
    run_uni = p_uni.add_run("TRƯỜNG CÔNG NGHỆ THÔNG TIN PHENIKAA\nĐẠI HỌC PHENIKAA\n")
    apply_run_font(run_uni, size_pt=14, bold=True)
    
    p_sep = doc.add_paragraph()
    p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sep = p_sep.add_run("-------------------- *** --------------------")
    apply_run_font(run_sep, size_pt=12, bold=True)
    
    # Khoảng trống giữa
    for _ in range(5):
        doc.add_paragraph()
        
    # Tên môn học
    p_subject = doc.add_paragraph()
    p_subject.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_subject.paragraph_format.space_after = Pt(12)
    run_subj = p_subject.add_run("BÁO CÁO BÀI TẬP GIỮA KỲ\nMÔN: ỨNG DỤNG PHÂN TÁN")
    apply_run_font(run_subj, size_pt=16, bold=True)
    
    # Tên đề tài
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(20)
    p_title.paragraph_format.space_after = Pt(40)
    run_title = p_title.add_run("ĐỀ TÀI: XÂY DỰNG HỆ THỐNG LƯU TRỮ KEY–VALUE PHÂN TÁN (DISTRIBUTED KEY-VALUE STORE)\nHỖ TRỢ CONSISTENT HASHING & REPLICATION")
    apply_run_font(run_title, size_pt=18, bold=True, color_rgb=(0, 51, 102)) # Màu xanh đậm sang trọng
    
    # Khoảng trống trước bảng thông tin
    for _ in range(4):
        doc.add_paragraph()

    # Bảng thông tin sinh viên và giáo viên
    p_info_title = doc.add_paragraph()
    p_info_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_info_title.paragraph_format.left_indent = Inches(1.5)
    r_info = p_info_title.add_run("Giảng viên hướng dẫn:  TS. Đỗ Quốc Trường\nLớp tín chỉ:                      N03\n\nDanh sách sinh viên thực hiện:")
    apply_run_font(r_info, size_pt=13, bold=True)

    # Tạo bảng danh sách sinh viên
    from docx.enum.table import WD_TABLE_ALIGNMENT
    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Cấu hình đường viền bảng mỏng/gọn gàng bằng cách set XML
    tblPr = table._tbl.tblPr
    borders = parse_xml(r'<w:tblBorders %s><w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/><w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/><w:left w:val="none"/><w:right w:val="none"/><w:insideV w:val="none"/></w:tblBorders>' % nsdecls('w'))
    tblPr.append(borders)

    headers = ["Họ và tên", "MSSV", "Lớp"]
    student_data = [
        ["Nguyễn Thị Quỳnh Anh", "23010147", "K17_CNTT2"],
        ["Nguyễn Thị Huyền Trang", "23010181", "K17_CNTT2"],
        ["Vương Đức Việt", "23010589", "K17_CNTT6"]
    ]

    # Điền header
    for i, head in enumerate(headers):
        cell = table.cell(0, i)
        cell.paragraphs[0].text = head
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        apply_run_font(cell.paragraphs[0].runs[0], size_pt=12, bold=True)
        # Background header
        shading_elm = parse_xml(r'<w:shd {} w:fill="F2F2F2"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading_elm)

    # Điền dữ liệu
    for row_idx, data in enumerate(student_data):
        for col_idx, val in enumerate(data):
            cell = table.cell(row_idx + 1, col_idx)
            cell.paragraphs[0].text = val
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            apply_run_font(cell.paragraphs[0].runs[0], size_pt=12)

    # Căn chỉnh kích thước cột rộng ra một chút
    table.columns[0].width = Inches(2.5)
    table.columns[1].width = Inches(1.5)
    table.columns[2].width = Inches(1.5)

    # Ngày tháng
    for _ in range(4):
        doc.add_paragraph()
        
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_date = p_date.add_run("Hà Nội, tháng 3 năm 2026")
    apply_run_font(r_date, size_pt=12, italic=True)

    # Ngắt trang
    doc.add_page_break()

    # ────────────────────────────────────────────────────────
    # 2. MỤC LỤC
    # ────────────────────────────────────────────────────────
    add_heading_1("Mục Lục")
    
    sections_list = [
        ["1. Giới thiệu", "3"],
        ["   1.1. Bài toán đặt ra", "3"],
        ["   1.2. Mục tiêu của hệ thống", "3"],
        ["2. Kiến trúc tổng thể hệ thống", "4"],
        ["   2.1. Mô hình kiến trúc phân tán", "4"],
        ["   2.2. Thành phần chi tiết của mỗi Node", "4"],
        ["   2.3. Giao thức truyền thông XML-RPC", "4"],
        ["   2.4. Cấu hình và khởi tạo cụm (Cluster Configuration)", "5"],
        ["3. Mô hình lưu trữ, phân mảnh dữ liệu và sao lưu", "5"],
        ["   3.1. Phân chia dữ liệu dựa trên thuật toán Consistent Hashing", "5"],
        ["   3.2. Mô hình sao lưu dữ liệu (Replication Model)", "6"],
        ["   3.3. Mô hình nhất quán dữ liệu (Eventual Consistency)", "6"],
        ["4. Định tuyến yêu cầu (Request Routing) và Chuyển tiếp (Forwarding)", "7"],
        ["   4.1. Định tuyến các yêu cầu ghi (PUT) và đọc (GET)", "7"],
        ["   4.2. Khả năng dự phòng lỗi khi định tuyến (Failover Routing)", "7"],
        ["5. Chiến lược phát hiện lỗi và xử lý node hỏng", "8"],
        ["   5.1. Cơ chế Heartbeat và phát hiện lỗi tự động", "8"],
        ["   5.2. Kịch bản khi có node bị sập (Node Failure Handling)", "8"],
        ["6. Cơ chế khôi phục và đồng bộ lại dữ liệu", "9"],
        ["   6.1. Đồng bộ và lọc phân vùng khi khởi động lại", "9"],
        ["   6.2. Đánh giá ưu và nhược điểm của cơ chế đồng bộ", "9"],
        ["7. Phân tích hệ thống theo định lý CAP", "10"],
        ["8. Ứng dụng quản trị trực quan (Flask Web Dashboard)", "10"],
        ["   8.1. Thiết kế giao diện và kiến trúc ứng dụng Web", "10"],
        ["   8.2. Các API Quản trị hệ thống phân tán", "11"],
        ["9. Đánh giá hạn chế và Đề xuất cải tiến", "11"],
        ["   9.1. Về xử lý đa luồng đồng thời và đồng bộ hóa (Thread-Safety)", "11"],
        ["   9.2. Hạn chế về cơ chế đồng bộ hóa ban đầu", "12"],
        ["   9.3. Đề xuất về lưu trữ bền vững (Persistence Storage)", "12"],
        ["10. Kết luận", "12"]
    ]

    for title, page in sections_list:
        p_toc = doc.add_paragraph()
        p_toc.paragraph_format.space_after = Pt(3)
        # Sử dụng tab stops hoặc dấu chấm chấm dẫn hướng đơn giản
        dots_count = 100 - len(title)*2
        dots = "." * max(10, dots_count)
        
        run_title = p_toc.add_run(title)
        apply_run_font(run_title, size_pt=12, bold=("   " not in title))
        
        run_dots = p_toc.add_run(f" {dots} ")
        apply_run_font(run_dots, size_pt=10, color_rgb=(150, 150, 150))
        
        run_page = p_toc.add_run(page)
        apply_run_font(run_page, size_pt=12, bold=("   " not in title))

    doc.add_page_break()

    # ────────────────────────────────────────────────────────
    # 3. NỘI DUNG CHÍNH
    # ────────────────────────────────────────────────────────
    
    # --- PHẦN 1 ---
    add_heading_1("1. Giới thiệu")
    
    add_heading_2("1.1. Bài toán đặt ra")
    add_body("Trong các hệ thống lưu trữ tập trung truyền thống, toàn bộ thông tin được ghi nhận và truy vấn thông qua một máy chủ (server) duy nhất. Cấu trúc này bộc lộ những nhược điểm nghiêm trọng khi ứng dụng mở rộng quy mô:")
    add_bullet("Điểm lỗi duy nhất (Single Point of Failure - SPOF): Khi máy chủ gặp sự cố vật lý hoặc sập mạng, toàn bộ dịch vụ sẽ ngừng hoạt động.", bold_prefix="o ")
    add_bullet("Giới hạn khả năng mở rộng (Scalability): Việc nâng cấp phần cứng (mở rộng theo chiều dọc - Vertical Scaling) có chi phí rất đắt đỏ và bị giới hạn bởi ngưỡng vật lý của thiết bị.", bold_prefix="o ")
    add_bullet("Hiệu năng nghẽn cổ chai: Khi lưu lượng truy cập tăng vọt, một máy chủ không thể đáp ứng đồng thời hàng nghìn kết nối đọc/ghi cùng lúc.", bold_prefix="o ")
    add_body("Để giải quyết triệt để các vấn đề trên, nhóm phát triển đã xây dựng giải pháp \"Hệ thống lưu trữ Key-Value phân tán\". Bằng cách kết nối nhiều máy chủ độc lập (các Node) hoạt động cùng nhau tạo thành một cụm (Cluster), hệ thống có khả năng chia sẻ tải trọng, tự động sao lưu dữ liệu và duy trì hoạt động liên tục ngay cả khi một số thành phần trong hệ thống bị hỏng.")

    add_heading_2("1.2. Mục tiêu của hệ thống")
    add_body("Mục tiêu chính của đề tài bao gồm:")
    add_bullet("Triển khai tối thiểu cụm 3 node độc lập, chia sẻ công việc lưu trữ dữ liệu.", bold_prefix="- ")
    add_bullet("Hỗ trợ đầy đủ 3 API thao tác cơ bản: PUT(key, value) để lưu trữ, GET(key) để truy vấn và DELETE(key) để xóa dữ liệu.", bold_prefix="- ")
    add_bullet("Phân mảnh dữ liệu thông minh nhằm chia sẻ đều dung lượng lưu trữ giữa các node.", bold_prefix="- ")
    add_bullet("Sao lưu (Replicate) mỗi bản ghi sang node lân cận để đảm bảo an toàn dữ liệu.", bold_prefix="- ")
    add_bullet("Xây dựng cơ chế Heartbeat để tự phát hiện khi có node gặp sự cố.", bold_prefix="- ")
    add_bullet("Tự động phục hồi và đồng bộ hóa lại dữ liệu một cách thông minh khi một node lỗi khởi động lại.", bold_prefix="- ")
    add_bullet("Cung cấp giao diện Web UI trực quan để quản trị, theo dõi trạng thái sống/chết của cụm và kiểm tra phân bổ dữ liệu.", bold_prefix="- ")

    # --- PHẦN 2 ---
    add_heading_1("2. Kiến trúc tổng thể hệ thống")
    
    add_heading_2("2.1. Mô hình kiến trúc phân tán")
    add_body("Hệ thống được thiết kế theo kiến trúc phi tập trung (Peer-to-Peer). Không có node master đóng vai trò điều phối trung tâm; thay vào đó, toàn bộ các node trong cluster đều có vai trò ngang hàng nhau. Mỗi node vừa có khả năng lưu trữ dữ liệu, vừa có khả năng tiếp nhận và xử lý yêu cầu trực tiếp từ client. Dữ liệu được phân bổ và sao lưu dựa trên vị trí của các node trong mạng băm nhất quán.")
    add_body("Các node giao tiếp song phương và đồng đẳng với nhau. Khách hàng (Client) có thể gửi yêu cầu đọc/ghi tới bất kỳ node nào trong cluster, node tiếp nhận sẽ tự động tính toán định tuyến để chuyển tiếp yêu cầu đến đúng vị trí lưu trữ thực tế.")

    add_heading_2("2.2. Thành phần chi tiết của mỗi Node")
    add_body("Mỗi node trong hệ thống được định nghĩa bởi một lớp đối tượng độc lập (`KeyValueNode`) chứa các thành phần cốt lõi:")
    add_bullet("data_store (Primary Store): Bộ nhớ lưu trữ chính (dạng dictionary trong RAM) chứa các cặp Key-Value mà node đó chịu trách nhiệm trực tiếp.", bold_prefix="1. ")
    add_bullet("replica_store (Replica Store): Bộ nhớ lưu trữ phụ dùng để chứa các bản sao (backup) của các key thuộc quyền quản lý của node lân cận.", bold_prefix="2. ")
    add_bullet("neighbors: Danh sách các cổng kết nối (port) và địa chỉ của các node lân cận trong cụm.", bold_prefix="3. ")
    add_bullet("Cơ chế Heartbeat: Bao gồm vòng lặp gửi tín hiệu ping định kỳ sang tất cả các neighbor và vòng lặp giám sát timeout nhằm cập nhật trạng thái sống/chết (ALIVE/DEAD) của các node lân cận.", bold_prefix="4. ")
    add_bullet("Cơ chế Lock đồng bộ: Sử dụng đối tượng Lock của thư viện `threading` để bảo vệ tài nguyên bộ nhớ tránh hiện tượng Race Condition do xử lý đồng thời nhiều truy vấn XML-RPC từ client hoặc node khác gửi đến.", bold_prefix="5. ")

    add_heading_2("2.3. Giao thức truyền thông XML-RPC")
    add_body("Mọi hoạt động giao tiếp truyền thông liên kết giữa Client và Node, cũng như giữa Node với Node đều được thực hiện thông qua giao thức XML-RPC (XML Remote Procedure Call) chạy trên nền HTTP và lớp mạng TCP:")
    add_bullet("Đặc tính: Đơn giản, thư viện chuẩn có sẵn trong Python (`xmlrpc.client`, `xmlrpc.server.SimpleXMLRPCServer`), giúp giảm thiểu thời gian cài đặt.", bold_prefix="- ")
    add_bullet("Tuần tự hóa (Serialization): Dữ liệu Python được tự động đóng gói dưới dạng XML truyền qua kết nối HTTP và được giải tuần tự hóa tự động ở phía nhận.", bold_prefix="- ")
    add_bullet("Tối ưu hóa Timeout: Sử dụng lớp `TimeoutTransport` tùy biến kế thừa từ lớp `xmlrpc.client.Transport` để áp đặt giới hạn thời gian chờ kết nối (1.5 giây). Điều này ngăn chặn việc các node bị treo (block) vĩnh viễn khi kết nối tới các node đã sập.", bold_prefix="- ")

    add_heading_2("2.4. Cấu hình và khởi tạo cụm (Cluster Configuration)")
    add_body("Hệ thống khởi chạy bằng cách cung cấp cổng lắng nghe riêng của node và danh sách toàn bộ các cổng cấu thành nên cluster thông qua đối số dòng lệnh:")
    add_body("python node.py [PORT_CUA_TOI] [PORT_CLUSTER_1] [PORT_CLUSTER_2] ...", italic=True, indent_inches=0.5)
    add_body("Ví dụ, để khởi động cụm 3 node chạy trên máy cục bộ (localhost) tại các cổng 8000, 8001 và 8002, ta sử dụng tệp lệnh `start_cluster.bat` thực thi các câu lệnh sau song song:")
    add_bullet("python node.py 8000 8000 8001 8002 (Node 1)", bold_prefix="* ")
    add_bullet("python node.py 8001 8000 8001 8002 (Node 2)", bold_prefix="* ")
    add_bullet("python node.py 8002 8000 8001 8002 (Node 3)", bold_prefix="* ")

    # --- PHẦN 3 ---
    add_heading_1("3. Mô hình lưu trữ, phân mảnh dữ liệu và sao lưu")
    
    add_heading_2("3.1. Phân chia dữ liệu dựa trên thuật toán Consistent Hashing")
    add_body("Khác biệt hoàn toàn với mô hình sao lưu toàn phần (nơi mọi node đều lưu tất cả dữ liệu), hệ thống thực tế triển khai giải pháp Phân chia phân vùng (Sharding/Partitioning) thông minh dựa trên kỹ thuật băm nhất quán (Consistent Hashing) đơn giản hóa:")
    add_body("Khi một key được tạo ra, hệ thống áp dụng hàm băm MD5 để chuyển đổi chuỗi ký tự key thành một số nguyên lớn, sau đó chia lấy dư cho tổng số node có trong cụm để tìm ra vị trí chỉ mục (index) của node Primary:")
    add_body("Index = int(MD5(Key), 16) % Tổng_Số_Node_Trong_Cluster", italic=True, indent_inches=0.5)
    add_body("Vị trí kết quả này ánh xạ trực tiếp đến một phần tử trong danh sách các node được sắp xếp tăng dần theo số cổng. Node nằm tại vị trí Index này được gọi là Primary Node chịu trách nhiệm lưu trữ chính cặp dữ liệu đó trong bộ nhớ `data_store`.")
    add_image("diagram_hash_ring.png", width_inches=4.0, caption="Hình 3.1. Sơ đồ Vòng tròn Consistent Hashing & Nhân bản bản sao")

    add_heading_2("3.2. Mô hình sao lưu dữ liệu (Replication Model)")
    add_body("Để đảm bảo an toàn dữ liệu trước nguy cơ sập node, mỗi khóa ghi nhận không chỉ được lưu ở một node duy nhất mà được tự động nhân bản (replicate) tới một node phụ. Hệ thống áp dụng cấu hình N-Replica với N=1 (tổng cộng có 2 bản sao dữ liệu trong cụm):")
    add_bullet("Xác định vị trí Replica: Đối với một key, node Replica được tính là node nằm ở vị trí tiếp theo ngay sau node Primary trên vòng tròn băm:", bold_prefix="- ")
    add_body("Replica_Index = (Primary_Index + 1) % Tổng_Số_Node_Trong_Cluster", italic=True, indent_inches=0.5)
    add_bullet("Ghi nhận bản sao: Khi node Primary nhận được yêu cầu PUT(key, value) từ client, nó ghi dữ liệu vào `data_store` cục bộ, sau đó tạo một luồng chạy nền (`threading.Thread`) kết nối không đồng bộ tới node Replica để ghi dữ liệu này vào `replica_store` của node Replica đó.", bold_prefix="- ")
    add_bullet("Bỏ qua khi lỗi: Nếu node Replica đang ở trạng thái sập mạng (DOWN), tiến trình ghi nền sẽ bỏ qua kết nối lỗi để không làm nghẽn hoặc trì hoãn phản hồi tới client.", bold_prefix="- ")

    add_heading_2("3.3. Mô hình nhất quán dữ liệu (Eventual Consistency)")
    add_body("Hệ thống tuân thủ mô hình nhất quán cuối cùng (Eventual Consistency):")
    add_bullet("Ghi nhanh phản hồi: Node Primary sau khi ghi thành công dữ liệu vào RAM cục bộ sẽ trả về phản hồi thành công ngay lập tức cho client mà không bắt buộc phải chờ tiến trình replicate sang node Replica hoàn thành.", bold_prefix="1. ")
    add_bullet("Đồng bộ phi tuần tự: Hành động nhân bản sang node Replica chạy bất đồng bộ dưới nền. Trong một khoảng thời gian cực ngắn (vài phần mười giây), dữ liệu giữa node Primary và Replica có thể chưa đồng bộ hoàn toàn. Tuy nhiên, sau đó dữ liệu sẽ hội tụ và đạt trạng thái nhất quán.", bold_prefix="2. ")
    add_bullet("Tối ưu hóa: Lựa chọn này giúp hệ thống đạt độ trễ cực thấp cho thao tác PUT, tăng cường khả năng hoạt động liên tục (Availability) theo định lý CAP.", bold_prefix="3. ")

    # --- PHẦN 4 ---
    add_heading_1("4. Định tuyến yêu cầu (Request Routing) và Chuyển tiếp (Forwarding)")
    
    add_heading_2("4.1. Định tuyến các yêu cầu ghi (PUT) và đọc (GET)")
    add_body("Do client có thể kết nối ngẫu nhiên tới bất kỳ node nào trong cluster để thực hiện yêu cầu, hệ thống triển khai cơ chế Định tuyến và Chuyển tiếp (Request Routing/Forwarding) cục bộ tại mỗi node:")
    add_bullet("Đối với lệnh PUT(key, value): Node tiếp nhận sẽ tính toán node Primary cho key. Nếu chính nó là Primary, nó lưu vào `data_store` rồi nhân bản sang Replica. Nếu node khác là Primary, nó đóng vai trò trung gian, thực hiện lệnh gọi XML-RPC chuyển tiếp toàn bộ dữ liệu ghi sang cho node Primary tương ứng để xử lý tiếp.", bold_prefix="- ")
    add_bullet("Đối với lệnh GET(key): Node tiếp nhận kiểm tra xem key có sẵn trong RAM cục bộ của nó hay không (bao gồm cả trong vùng `data_store` chính hoặc `replica_store` sao lưu). Nếu có, nó trả kết quả ngay lập tức (Zero-hop). Nếu không có, nó sẽ tự động gửi yêu cầu XML-RPC chuyển tiếp tới node Primary của key đó.", bold_prefix="- ")

    add_heading_2("4.2. Khả năng dự phòng lỗi khi định tuyến (Failover Routing)")
    add_body("Khi xảy ra sự cố sập mạng hoặc một node trong cụm bị chết hoàn toàn, hệ thống vẫn đảm bảo định tuyến thông minh không gây lỗi kết nối cho client:")
    add_bullet("Trong lệnh PUT: Nếu node Primary của key đó bị sập, node trung gian khi gửi forward sẽ gặp ngoại lệ kết nối. Lúc này nó sẽ bắt ngoại lệ và tự động chuyển hướng (fallback) ghi dữ liệu trực tiếp vào node Replica. Dữ liệu tạm thời lưu trong vùng sao lưu của Replica.", bold_prefix="- ")
    add_bullet("Trong lệnh GET: Nếu node Primary bị sập, node nhận request sẽ phát hiện lỗi kết nối, sau đó tự động chuyển hướng kết nối sang node Replica để lấy bản sao dữ liệu và trả lại cho Client.", bold_prefix="- ")
    add_body("Nhờ vậy, hệ thống chịu đựng được việc sập một node bất kỳ mà không làm gián đoạn khả năng đọc ghi của các key đang hoạt động.")
    add_image("diagram_request_flow.png", width_inches=4.8, caption="Hình 4.1. Sơ đồ Luồng Định tuyến và Chuyển tiếp Request (Failover Routing)")

    # --- PHẦN 5 ---
    add_heading_1("5. Chiến lược phát hiện lỗi và xử lý node hỏng")
    
    add_heading_2("5.1. Cơ chế Heartbeat và phát hiện lỗi tự động")
    add_body("Mỗi node duy trì hai tiến trình chạy nền độc lập dưới dạng daemon thread để theo dõi sức khỏe của cụm:")
    add_bullet("Gửi Heartbeat (Vòng lặp gửi): Luồng `_send_heartbeat_loop` định kỳ cứ mỗi 3 giây sẽ thực hiện gọi hàm RPC `heartbeat(self.port)` tới toàn bộ các node neighbor có trong danh sách để thông báo nó vẫn đang hoạt động.", bold_prefix="1. ")
    add_bullet("Nhận và lưu dấu: Khi nhận được lệnh gọi RPC `heartbeat`, node nhận sẽ cập nhật thời gian nhận tin nhắn cuối cùng `self.last_heartbeat[from_port] = time.time()` và đánh dấu trạng thái node gửi là ALIVE.", bold_prefix="2. ")
    add_bullet("Giám sát lỗi (Vòng lặp quét): Luồng `_detect_failures_loop` định kỳ quét danh sách neighbor cứ mỗi 2 giây. Nếu phát hiện thời gian trôi qua kể từ lần cuối nhận heartbeat từ một node vượt quá 10 giây (TIMEOUT = 10), node giám sát sẽ chính thức đánh dấu node đó là DEAD.", bold_prefix="3. ")

    add_heading_2("5.2. Kịch bản khi có node bị sập (Node Failure Handling)")
    add_body("Giả sử cụm gồm 3 node A, B, C và node B bị dừng hoạt động:")
    add_bullet("Phát hiện lỗi: A và C sẽ nhận thấy tín hiệu heartbeat từ B dừng lại quá 10 giây, lập tức cập nhật trạng thái B thành DEAD trong bảng trạng thái nội bộ của mình.", bold_prefix="- ")
    add_bullet("Bỏ qua truyền tải lỗi: Trong các tiến trình replicate dữ liệu của A và C tiếp theo, các luồng gửi phụ sẽ phát hiện B là DEAD hoặc gặp lỗi timeout nhanh 1.5 giây nên sẽ bỏ qua ngay lập tức, không gây treo luồng chính xử lý request của client.", bold_prefix="- ")
    add_bullet("Client vẫn kết nối an toàn: Yêu cầu ghi/đọc của các key thuộc quản lý của B sẽ được tự động định tuyến đến node replica của nó (ví dụ C) để thực thi.", bold_prefix="- ")

    # --- PHẦN 6 ---
    add_heading_1("6. Cơ chế khôi phục và đồng bộ lại dữ liệu")
    
    add_heading_2("6.1. Đồng bộ và lọc phân vùng khi khởi động lại")
    add_body("Khi một node bị lỗi được khởi chạy lại, nó cần khôi phục lại trạng thái dữ liệu đã mất đi trong bộ nhớ RAM của mình. Cơ chế phục hồi được mã nguồn xây dựng thông qua hàm `_sync_data_on_startup()`:")
    add_bullet("Kết nối neighbor: Node mới khởi động sẽ duyệt qua danh sách các cổng của các neighbor. Nó cố gắng kết nối tới node neighbor đầu tiên đang hoạt động.", bold_prefix="1. ")
    add_bullet("Tải snapshot thô: Gọi hàm RPC `get_all_data()` để nhận toàn bộ dữ liệu lưu trữ (gồm cả Primary và Replica) của node neighbor đó dưới định dạng chuỗi JSON.", bold_prefix="2. ")
    add_bullet("Tính toán băm và lọc khóa (Điểm cốt lõi): Sau khi giải mã dữ liệu thô, node khởi động duyệt qua từng key nhận được. Nó gọi hàm tính toán băm nhất quán `_get_primary_port(key)` và `_get_replica_port(key)`. Node mới chỉ nạp key đó vào bộ nhớ của mình nếu chính nó là Primary (lưu vào `data_store`) hoặc chính nó là Replica (lưu vào `replica_store`). Các key không thuộc phân vùng quản lý của nó sẽ bị bỏ qua hoàn toàn.", bold_prefix="3. ")
    add_bullet("Hoàn thành: Trạng thái dữ liệu của node được tái cấu trúc khớp chính xác với bản đồ phân vùng của cluster.", bold_prefix="4. ")
    add_image("diagram_heartbeat_sync.png", width_inches=6.0, caption="Hình 5.1. Sơ đồ Cơ chế Heartbeat và Đồng bộ dữ liệu Lọc Key")

    add_heading_2("6.2. Đánh giá ưu và nhược điểm của cơ chế đồng bộ")
    add_body("Ưu điểm:", bold_prefix="* ")
    add_bullet("Thiết kế thông minh hơn đồng bộ toàn phần: Dữ liệu sau khi đồng bộ được phân vùng chính xác về các node tương ứng, không gây dư thừa lưu trữ.", indent_inches=0.5)
    add_bullet("Tính tự động cao: Không yêu cầu can thiệp thủ công từ quản trị viên, node tự phục hồi ngay khi khởi động.", indent_inches=0.5)
    add_body("Nhược điểm:", bold_prefix="* ")
    add_bullet("Lấy dữ liệu thô lớn: Vẫn phải tải toàn bộ tập key-value thô từ một node khác về để phân tích, gây tốn băng thông đường truyền nếu cơ sở dữ liệu lớn.", indent_inches=0.5)
    add_bullet("Phụ thuộc tính khả dụng của node khác: Nếu tất cả các node khác đều sập vào thời điểm khởi động, node mới sẽ không thể phục hồi dữ liệu cũ từ bộ nhớ RAM của các node trước đó.", indent_inches=0.5)

    # --- PHẦN 7 ---
    add_heading_1("7. Phân tích hệ thống theo định lý CAP")
    add_body("Định lý CAP chỉ ra rằng một hệ thống phân tán chỉ có thể đáp ứng tối đa 2 trong số 3 yếu tố: Consistency (Tính nhất quán), Availability (Tính sẵn sàng) và Partition Tolerance (Tính chịu lỗi phân vùng).")
    add_body("Hệ thống Key-Value Store phân tán thực tế lựa chọn thiết kế ưu tiên cặp AP (Availability và Partition Tolerance):")
    add_bullet("Availability (Tính sẵn sàng cao): Mỗi node trong cluster sẵn sàng phản hồi nhanh yêu cầu PUT/GET của client mà không cần chờ sự đồng ý từ toàn bộ các node khác. Hệ thống sử dụng timeout 1.5 giây để tránh chặn tiến trình.", bold_prefix="1. ")
    add_bullet("Partition Tolerance (Chịu lỗi phân vùng): Nhờ cơ chế phát hiện lỗi Heartbeat và khả năng định tuyến chuyển hướng Failover sang node Replica, hệ thống vẫn duy trì dịch vụ đọc ghi ổn định ngay cả khi mạng bị chia cắt hoặc một số node bị sập hoàn toàn.", bold_prefix="2. ")
    add_bullet("Đánh đổi Consistency (Tính nhất quán mạnh): Hệ thống không đảm bảo tính nhất quán tức thời (Strong Consistency) do cơ chế replicate phi tuần tự và bất đồng bộ dưới nền. Dữ liệu giữa các node có thể lệch nhau trong khoảnh khắc ngắn trước khi hội tụ về trạng thái nhất quán cuối cùng (Eventual Consistency).", bold_prefix="3. ")

    # --- PHẦN 8 ---
    add_heading_1("8. Ứng dụng quản trị trực quan (Flask Web Dashboard)")
    
    add_heading_2("8.1. Thiết kế giao diện và kiến trúc ứng dụng Web")
    add_body("Để nâng cao khả năng quản trị và giám sát hệ thống phân tán trực quan, một ứng dụng Web Dashboard đi kèm đã được xây dựng bằng Flask framework (`manager_app.py`). Ứng dụng này giao tiếp với các node trong cluster bằng cách đóng vai trò là một client XML-RPC để thu thập dữ liệu và hiển thị lên giao diện HTML/CSS hiện đại.")
    add_body("Giao diện quản lý bao gồm các màn hình chính:")
    add_bullet("Trang Dashboard: Hiển thị các chỉ số tổng quan như Sức khỏe cụm (Cluster Health) tính bằng tỷ lệ phần trăm node hoạt động, số lượng node ONLINE, tổng số lượng key chính (Primary Keys) và key sao lưu (Replica Keys) trên toàn cụm.", bold_prefix="- ")
    add_bullet("Trang Quản lý Node (Server list): Liệt kê danh sách chi tiết các node hoạt động, hiển thị danh sách neighbor của từng node và trạng thái sống chết của chúng dưới góc nhìn của node đó.", bold_prefix="- ")
    add_bullet("Trang Quản lý Dữ liệu (Records): Cho phép xem danh sách tất cả các Key-Value đang tồn tại trong hệ thống, chỉ rõ key đó đang nằm ở node nào làm Primary và được sao lưu ở node nào làm Replica.", bold_prefix="- ")
    add_bullet("Trang Thử nghiệm ghi/đọc (PUT/GET/DELETE): Cung cấp các biểu mẫu nhập liệu giúp kiểm tra trực tiếp khả năng ghi nhận dữ liệu thông qua giao diện Web.", bold_prefix="- ")

    add_heading_2("8.2. Các API Quản trị hệ thống phân tán")
    add_body("Ứng dụng quản trị Flask cung cấp các API JSON phục vụ giao diện động:")
    add_bullet("/api/cluster/status: Trả về chi tiết cấu hình và trạng thái kết nối thực tế của từng node trong cụm.", bold_prefix="- ")
    add_bullet("/api/all-data: Quét toàn bộ dữ liệu từ tất cả các node trong cụm, gộp thông tin lại để chỉ ra vị trí chính và phụ của từng key phục vụ hiển thị bảng dữ liệu.", bold_prefix="- ")
    add_bullet("/api/routing/<key>: Trả về thông tin định tuyến cho một key cụ thể (nút nào quản lý chính, nút nào sao lưu).", bold_prefix="- ")
    add_bullet("/api/sync/<port>: Cho phép kích hoạt thủ công lệnh đồng bộ cưỡng bức (`force_sync`) trên một node để buộc nó dọn sạch dữ liệu cũ và tiến hành đồng bộ lại từ đầu với cụm.", bold_prefix="- ")

    # --- PHẦN 9 ---
    add_heading_1("9. Đánh giá hạn chế và Đề xuất cải tiến")
    
    add_heading_2("9.1. Về xử lý đa luồng đồng thời và đồng bộ hóa (Thread-Safety)")
    add_body("Khác với nhận định ban đầu cho rằng hệ thống thiếu kiểm soát truy cập đồng thời dẫn tới race condition, mã nguồn thực tế đã triển khai giải pháp sử dụng khóa loại trừ tương hỗ `threading.Lock()` bảo vệ an toàn cho các cấu trúc dữ liệu:")
    add_body("Cụ thể, đối tượng `self.lock` được dùng để khóa tài nguyên trong toàn bộ các khối lệnh thay đổi giá trị như `put`, `get`, `delete`, `heartbeat` và `get_node_info`. Thiết kế này đảm bảo an toàn tuyệt đối cho bộ nhớ RAM của các node khi chịu tải lớn từ nhiều luồng kết nối XML-RPC đồng thời.")
    add_body("Đề xuất cải tiến tiếp theo: Thay thế khóa toàn cục đơn giản bằng khóa phân mảnh (Read-Write Lock) để cho phép nhiều luồng cùng đọc dữ liệu đồng thời, chỉ khóa khi có thao tác ghi dữ liệu, giúp tăng tốc độ đáp ứng của hệ thống.")

    add_heading_2("9.2. Hạn chế về cơ chế đồng bộ hóa ban đầu")
    add_body("Hạn chế:")
    add_bullet("Khi đồng bộ dữ liệu lúc khởi động, node mới lấy toàn bộ dữ liệu từ node neighbor đầu tiên nó kết nối được. Nếu node neighbor này cũng vừa mới được khởi động và chưa kịp cập nhật đầy đủ dữ liệu từ các node khác, node mới sẽ bị thiếu dữ liệu nghiêm trọng.", bold_prefix="- ")
    add_bullet("Chưa có cơ chế so sánh phiên bản (versioning) hoặc nhãn thời gian (timestamps) cho các key, dễ dẫn đến hiện tượng dữ liệu cũ ghi đè lên dữ liệu mới hơn trong quá trình khôi phục.", bold_prefix="- ")
    add_body("Đề xuất cải tiến:")
    add_bullet("Triển khai cơ chế đồng bộ từ nhiều node khác nhau để đối chiếu dữ liệu chéo.", bold_prefix="- ")
    add_bullet("Tích hợp Vector Clocks hoặc nhãn thời gian vật lý gắn liền với mỗi giá trị key-value để giải quyết xung đột ghi đè dữ liệu cũ.", bold_prefix="- ")

    add_heading_2("9.3. Đề xuất về lưu trữ bền vững (Persistence Storage)")
    add_body("Hạn chế:")
    add_bullet("Hiện tại dữ liệu được lưu hoàn toàn trong bộ nhớ RAM (`data_store = {}`). Khi toàn bộ cụm node bị tắt điện đột ngột hoặc khởi động lại cùng lúc, toàn bộ cơ sở dữ liệu sẽ biến mất hoàn toàn.", bold_prefix="- ")
    add_body("Đề xuất cải tiến:")
    add_bullet("Bổ sung cơ chế lưu trữ dữ liệu xuống ổ đĩa cứng (Disk Persistence) bằng cách định kỳ ghi snapshot ra tệp JSON hoặc sử dụng cơ sở dữ liệu gọn nhẹ SQLite.", bold_prefix="- ")
    add_bullet("Áp dụng cơ chế Write-Ahead Logging (WAL) để ghi nhận nhật ký mọi thay đổi trước khi thực thi vào RAM, giúp phục hồi dữ liệu chính xác về thời điểm trước khi sập nguồn.", bold_prefix="- ")

    # --- PHẦN 10 ---
    add_heading_1("10. Kết luận")
    add_body("Hệ thống lưu trữ Key-Value phân tán được xây dựng thành công đã đạt được toàn bộ yêu cầu đề ra của một bài tập lớn ứng dụng phân tán học thuật:")
    add_bullet("Thiết kế thành công cụm nhiều node hoạt động song song, không phụ thuộc vào node trung tâm.", bold_prefix="- ")
    add_bullet("Cài đặt hiệu quả thuật toán Consistent Hashing để phân hoạch dữ liệu và cơ chế nhân bản bản sao đảm bảo khả năng chịu lỗi tốt.", bold_prefix="- ")
    add_bullet("Hiện thực hóa cơ chế phát hiện lỗi dựa trên giao thức Heartbeat và tự động chuyển hướng định tuyến khi có node chết mạng.", bold_prefix="- ")
    add_bullet("Giao diện Web UI giám sát hiện đại, giúp dễ dàng theo dõi trực quan và vận hành thử nghiệm cụm phân tán.", bold_prefix="- ")
    add_body("Hệ thống hoạt động ổn định trong môi trường mô phỏng cục bộ và là nền tảng vững chắc để tiếp tục nghiên cứu, tích hợp các thuật toán đồng thuận phức tạp hơn như Raft hoặc Paxos trong tương lai.")

    # Lưu file
    output_filename = "BaoCao_Moi.docx"
    try:
        doc.save(output_filename)
        print(f"File '{output_filename}' generated successfully!")
    except PermissionError:
        output_filename = "BaoCao_Moi_Sodo.docx"
        doc.save(output_filename)
        print(f"File '{output_filename}' generated successfully (fallback due to lock)!")

if __name__ == "__main__":
    create_report()
